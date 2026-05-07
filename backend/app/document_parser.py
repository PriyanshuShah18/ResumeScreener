from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path

from PIL import Image

from app.config import Settings
from app.constants import SUPPORTED_EXTENSIONS
from app.ocr import OCRUnavailableError, annotate_resume_sections, get_structured_ocr
from app.preprocess import preprocess

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from pdf2image import convert_from_bytes
except ImportError:  # pragma: no cover
    convert_from_bytes = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None


@dataclass
class ParsedDocument:
    text: str
    page_images: list[Image.Image] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    page_count: int = 0


class DocumentParser:
    def __init__(self, settings: Settings):
        self.settings = settings

    def is_supported(self, filename: str) -> bool:
        return Path(filename or "").suffix.lower() in SUPPORTED_EXTENSIONS

    def parse_upload(self, filename: str, content: bytes) -> ParsedDocument:
        if not self.is_supported(filename):
            raise ValueError(f"Unsupported file type for '{filename}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

        max_bytes = self.settings.max_file_size_mb * 1024 * 1024
        if len(content) > max_bytes:
            raise ValueError(f"File '{filename}' exceeds the {self.settings.max_file_size_mb}MB limit")

        extension = Path(filename).suffix.lower()
        if extension == ".pdf":
            return self._parse_pdf(content)
        if extension == ".docx":
            return self._parse_docx(content)
        return self._parse_image(content)

    def _parse_pdf(self, content: bytes, force_vision: bool = True) -> ParsedDocument:
        warnings: list[str] = []
        page_count = 0

        if PdfReader:
            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            selected_pages = reader.pages[: self.settings.max_pages_per_resume]
            if page_count > self.settings.max_pages_per_resume:
                warnings.append(
                    f"Processed only the first {self.settings.max_pages_per_resume} pages out of {page_count}"
                )
            direct_text = "\n".join((page.extract_text() or "") for page in selected_pages).strip()
            # If text is good enough, return immediately
            if len(re.sub(r"\s+", "", direct_text)) >= self.settings.min_pdf_text_chars:
                return ParsedDocument(
                    text=annotate_resume_sections(direct_text),
                    warnings=warnings,
                    page_count=page_count,
                )

        if not convert_from_bytes:
            raise RuntimeError("PDF image conversion is unavailable because pdf2image is not installed")

        convert_kwargs = {"first_page": 1, "last_page": self.settings.max_pages_per_resume}
        if self.settings.poppler_path:
            convert_kwargs["poppler_path"] = self.settings.poppler_path

        # Convert PDF to images for VLM/OCR
        images = convert_from_bytes(content, **convert_kwargs)
        if not images:
            raise ValueError("Unable to read PDF pages")

        if not page_count:
            page_count = len(images)

        page_images: list[Image.Image] = [img.convert("RGB") for img in images]

        # SPEED OPTIMIZATION: If force_vision is True, we skip Tesseract OCR entirely.
        # This saves 5-10 seconds per page and lets the VLM handle extraction from the images.
        if force_vision:
            return ParsedDocument(
                text="", # VLM will use images as primary source
                page_images=page_images,
                warnings=warnings,
                page_count=page_count,
            )

        # Legacy OCR Fallback (only if force_vision=False)
        page_texts: list[str] = []
        for rgb_image in page_images:
            try:
                page_texts.append(get_structured_ocr(preprocess(rgb_image)))
            except OCRUnavailableError as exc:
                warnings.append(f"OCR failed: {exc}. Using vision fallback.")
                break

        return ParsedDocument(
            text="\n\n".join(page_texts).strip(),
            page_images=page_images,
            warnings=warnings,
            page_count=page_count,
        )


    def _parse_docx(self, content: bytes) -> ParsedDocument:
        if not Document:
            raise RuntimeError("DOCX parsing is unavailable because python-docx is not installed")

        document = Document(io.BytesIO(content))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(row_text)

        text = "\n".join(chunks).strip()
        if not text:
            raise ValueError("DOCX file does not contain readable text")

        return ParsedDocument(text=annotate_resume_sections(text), page_count=1)

    def _parse_image(self, content: bytes) -> ParsedDocument:
        try:
            image = Image.open(io.BytesIO(content)).convert("RGB")
        except Exception as exc:  # pragma: no cover
            raise ValueError("Invalid image file") from exc

        try:
            structured_text = get_structured_ocr(preprocess(image))
        except OCRUnavailableError as exc:
            return ParsedDocument(
                text="",
                page_images=[image],
                warnings=[
                    f"OCR unavailable for this machine: {exc}. Falling back to vision-only extraction if the model is available."
                ],
                page_count=1,
            )

        if not structured_text:
            return ParsedDocument(
                text="",
                page_images=[image],
                warnings=[
                    "OCR could not recover readable text from the image. Falling back to vision-only extraction if the model is available."
                ],
                page_count=1,
            )

        return ParsedDocument(text=structured_text, page_images=[image], page_count=1)
